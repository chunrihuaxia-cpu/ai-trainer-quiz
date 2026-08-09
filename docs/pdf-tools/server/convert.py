"""PDF to Word / Image converter — with table support"""
import sys, os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import fitz

def to_word(in_path, out_path):
    doc = fitz.open(in_path)
    d = Document()

    # Default font
    style = d.styles['Normal']
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.insert(0, rFonts)
    style.font.size = Pt(10.5)

    for page_idx, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        page_w = page.rect.width

        # Detect tables
        table_data = []
        table_bounds = []
        try:
            tf = page.find_tables()
            if tf and tf.tables:
                for tab in tf.tables:
                    data = tab.extract()
                    if data:
                        table_data.append(data)
                        table_bounds.append(tab.bbox)
        except Exception:
            pass

        def in_table(x, y):
            return any(tx0 <= x <= tx1 and ty0 <= y <= ty1 for (tx0, ty0, tx1, ty1) in table_bounds)

        # Write text blocks
        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                spans = line["spans"]
                if not spans:
                    continue
                text = "".join(s["text"] for s in spans).strip()
                if not text:
                    continue
                # Skip table text (written separately as Word tables)
                bx, by = line["bbox"][0], line["bbox"][1]
                if in_table(bx, by):
                    continue

                first = spans[0]
                size_pt = first["size"] if first["size"] > 0 else 10.5
                is_bold = bool(first["flags"] & 16)
                x0 = line["bbox"][0]
                w = line["bbox"][2] - x0
                center_x = x0 + w / 2

                if abs(center_x - page_w / 2) < page_w * 0.12 and size_pt >= 16:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    align = WD_ALIGN_PARAGRAPH.LEFT

                p = d.add_paragraph()
                p.alignment = align
                run = p.add_run(text)
                run.font.size = Pt(max(size_pt, 8))
                run.bold = is_bold

        # Write detected tables
        for data in table_data:
            d.add_paragraph()
            rows = len(data)
            cols = max(len(row) for row in data) if data else 1
            wt = d.add_table(rows=rows, cols=cols, style='Table Grid')
            for i, row in enumerate(data):
                for j, cell_text in enumerate(row):
                    if j < cols:
                        cell = wt.cell(i, j)
                        cell.text = str(cell_text or '')
                        for p2 in cell.paragraphs:
                            for r2 in p2.runs:
                                r2.font.size = Pt(9)
            d.add_paragraph()

        if page_idx < len(doc) - 1:
            d.add_page_break()

    doc.close()
    d.save(out_path)
    print("OK")

def to_image(in_path, out_dir):
    doc = fitz.open(in_path)
    os.makedirs(out_dir, exist_ok=True)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        pix.save(f'{out_dir}/page_{i+1:03d}.png')
    doc.close()
    print("OK")

if __name__ == '__main__':
    cmd = sys.argv[1]
    if cmd == 'word':
        to_word(sys.argv[2], sys.argv[3])
    elif cmd == 'image':
        to_image(sys.argv[2], sys.argv[3])
